import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import zipfile
from datetime import datetime

# Configuração da página
st.set_page_config(page_title="Gerador Automático RSC", layout="wide")

st.title("Gerador de Declarações RSC em Lote")
st.write("Faça o upload de uma planilha Excel contendo os dados dos servidores para gerar as declarações automaticamente.")

# Mapeamento de meses para a formatação da data
meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", 
         "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]

def gerar_documento(nome, cargo, siape, unidade, sistemas_selecionados):
    doc = Document()
    
    # Título
    titulo = doc.add_paragraph()
    run_titulo = titulo.add_run("DECLARAÇÃO")
    run_titulo.bold = True
    run_titulo.font.size = Pt(14)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("")
    
    # Primeiro parágrafo
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p1.add_run("Declaramos, para os devidos fins de comprovação junto ao processo de Reconhecimento de Saberes e Competências (RSC), em atendimento ao disposto no Anexo IV do Decreto nº 13.048, de 3 de julho de 2026, que o(a) servidor(a) ")
    
    p1.add_run(nome).bold = True
    p1.add_run(", ocupante do cargo de ")
    p1.add_run(cargo).bold = True
    p1.add_run(", matrícula SIAPE nº ")
    p1.add_run(siape).bold = True
    p1.add_run(", lotado(a) no(a) ")
    p1.add_run(unidade).bold = True
    p1.add_run(", possui ou possuiu o perfil de operador/executor desempenhando atividades que demandam acesso e utilização dos seguintes sistemas estruturantes do governo federal:")
    
    # Lista numérica de sistemas baseada no Excel
    if sistemas_selecionados:
        for i, sis in enumerate(sistemas_selecionados, 1):
            p_sis = doc.add_paragraph(f"{i})\t{sis}")
            p_sis.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p_vazio = doc.add_paragraph("Nenhum sistema estruturante registrado.")
        p_vazio.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Segundo parágrafo
    p2 = doc.add_paragraph("Esta declaração é a expressão da verdade e destina-se exclusivamente à instrução de processo administrativo de concessão da Gratificação por Reconhecimento de Saberes e Competências (GRSC).")
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph("")
    
    # Data alinhada à esquerda com ponto final
    hoje = datetime.now()
    data_atual = f"Belo Horizonte, {hoje.day} de {meses[hoje.month - 1]} de {hoje.year}."
    
    p_data = doc.add_paragraph(data_atual)
    p_data.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph("\n")
    
    # Assinatura
    assinatura = doc.add_paragraph()
    assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ass = assinatura.add_run("ELÍZIO MARCOS DOS REIS\n")
    run_ass.bold = True
    assinatura.add_run("Diretor do Departamento de Contabilidade e Finanças\n")
    assinatura.add_run("Pró-Reitoria de Planejamento e Desenvolvimento - UFMG")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def gerar_zip_declaracoes(df):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for index, row in df.iterrows():
            nome = str(row.get('nome', f'Servidor_{index}')).strip()
            
            # Flexibilidade: busca na coluna 'cargo', se não achar, tenta 'unidade'
            cargo = str(row.get('cargo', row.get('unidade', ''))).strip()
            # Leitura exclusiva da coluna SIAPE
            siape = str(row.get('siape', '')).strip()
            unidade_dept = str(row.get('uadnome', '')).strip()
            
            sistemas = []
            if str(row.get('siafi', '')).strip().upper() == 'X':
                sistemas.append("Sistema de Administração Financeira do Governo Federal (SIAFI);")
            if str(row.get('scdp', '')).strip().upper() == 'X':
                sistemas.append("Sistema Concessão de Diárias e Passagens (SCDP);")
            if str(row.get('tg', '')).strip().upper() == 'X':
                sistemas.append("Tesouro Gerencial (TG)")
                
            if sistemas:
                ultimo_item = sistemas[-1]
                if ultimo_item.endswith(';'):
                    sistemas[-1] = ultimo_item[:-1] + "."
                elif not ultimo_item.endswith('.'):
                    sistemas[-1] = ultimo_item + "."
            
            docx_buffer = gerar_documento(nome, cargo, siape, unidade_dept, sistemas)
            nome_arquivo = f"Declaracao_RSC_{nome.replace(' ', '_')}.docx"
            zip_file.writestr(nome_arquivo, docx_buffer.getvalue())
            
    zip_buffer.seek(0)
    return zip_buffer

st.divider()

st.subheader("Carregar Dados")
st.info("A planilha deve conter as colunas: **NOME**, **CARGO** (ou **UNIDADE**), **SIAPE**, **UADNOME**, **SOLICITADO** e as colunas de sistemas **SIAFI**, **SCDP** e **TG**.")

uploaded_file = st.file_uploader("Envie a planilha Excel (.xlsx)", type=["xlsx", "xls"])

if uploaded_file:
    # Tratando as células vazias lendo o Excel estritamente como string
    df = pd.read_excel(uploaded_file, dtype=str).fillna("")
    
    # Padronizando o nome das colunas
    df.columns = df.columns.str.strip().str.lower()
    
    # Nova lista de colunas obrigatórias com o 'siape' correto
    colunas_obrigatorias = ['nome', 'siape', 'uadnome', 'siafi', 'scdp', 'tg', 'solicitado']
    colunas_presentes = [col for col in colunas_obrigatorias if col in df.columns]
    
    # Verifica se há a coluna de cargo ou unidade
    tem_cargo = 'cargo' in df.columns or 'unidade' in df.columns
    
    if len(colunas_presentes) < len(colunas_obrigatorias) or not tem_cargo:
        faltantes = set(colunas_obrigatorias) - set(colunas_presentes)
        if not tem_cargo:
            faltantes.add("cargo (ou unidade)")
        st.error(f"Erro: A planilha enviada não contém todas as colunas esperadas. Colunas faltando: {', '.join(faltantes)}")
    else:
        # Filtra o DataFrame apenas para quem possui 'X' (maiúsculo ou minúsculo) na coluna 'solicitado'
        df_filtrado = df[df['solicitado'].str.strip().str.upper() == 'X']
        
        if len(df_filtrado) > 0:
            # LÓGICA DE BLOQUEIO: Verifica se todos os sistemas estão em branco para algum servidor solicitado
            cond_siafi = df_filtrado['siafi'].str.strip().str.upper() == 'X'
            cond_scdp = df_filtrado['scdp'].str.strip().str.upper() == 'X'
            cond_tg = df_filtrado['tg'].str.strip().str.upper() == 'X'
            
            df_sem_sistema = df_filtrado[~(cond_siafi | cond_scdp | cond_tg)]
            
            if not df_sem_sistema.empty:
                st.error("❌ Processamento Bloqueado: Foram encontrados servidores com a declaração solicitada, mas sem nenhum sistema marcado.")
                st.write("Verifique a lista abaixo, preencha as colunas SIAFI, SCDP ou TG na sua planilha e faça o upload novamente:")
                
                col_cargo_exibicao = 'cargo' if 'cargo' in df.columns else 'unidade'
                st.dataframe(df_sem_sistema[['nome', col_cargo_exibicao, 'siape', 'siafi', 'scdp', 'tg']])
            else:
                st.success(f"Planilha aprovada! {len(df_filtrado)} servidor(es) com declaração solicitada ('X') prontos para geração.")
                
                # Seleciona qual coluna exibir no preview da tabela
                col_cargo_exibicao = 'cargo' if 'cargo' in df.columns else 'unidade'
                st.dataframe(df_filtrado[['nome', col_cargo_exibicao, 'siape', 'uadnome', 'solicitado', 'siafi', 'scdp', 'tg']].head())
                
                if st.button("Gerar Declarações Solicitadas (Arquivo ZIP)", type="primary"):
                    with st.spinner("Lendo cruzamentos e gerando documentos..."):
                        zip_file = gerar_zip_declaracoes(df_filtrado)
                        
                    st.download_button(
                        label="📦 Baixar Lote Filtrado (ZIP)",
                        data=zip_file,
                        file_name="Declaracoes_RSC_Lote_Filtrado.zip",
                        mime="application/zip",
                        type="primary",
                        use_container_width=True
                    )
        else:
            st.warning("Nenhuma declaração foi marcada com 'X' na coluna 'SOLICITADO' para ser gerada.")
